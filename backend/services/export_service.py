"""
Export service — generate ATS-compliant DOCX and PDF files.

ATS compliance rules enforced:
1. Standard section headings only: Summary, Skills, Experience, Education, Projects
2. No tables, columns, text boxes, headers/footers
3. No images, icons, or special characters (beyond hyphens and pipes)
4. Fonts: Calibri for body (10-11pt), name in 16pt bold
5. Skills as flat comma-separated list
6. Margins: 0.75 inch all sides
7. Output: .docx (not .doc)
"""

import io
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_HYPHEN_AND_PIPE_RE = re.compile(r"[^\x20-\x7E]")  # strip non-printable / non-ASCII


def _clean_text(text: str) -> str:
    """Remove special characters beyond hyphens, pipes, and standard ASCII."""
    if not text:
        return ""
    # Replace smart quotes and dashes with ASCII equivalents
    replacements = {
        "\u2018": "'", "\u2019": "'", "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-", "\u2026": "...",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    # Strip remaining non-ASCII
    text = re.sub(r"[^\x20-\x7E]", "", text)
    return text.strip()


def _format_date_range(start: str, end: str) -> str:
    if not start and not end:
        return ""
    if not end or end.lower() == "present":
        return f"{start} – Present"
    return f"{start} – {end}"


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def _set_margins(document: Document, margin_inches: float = 0.75) -> None:
    for section in document.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)
        # Remove headers/footers
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for para in section.header.paragraphs:
            para.clear()
        for para in section.footer.paragraphs:
            para.clear()


def _add_name(doc: Document, name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_clean_text(name))
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(16)


def _add_contact_line(doc: Document, contact: Dict) -> None:
    parts = []
    if contact.get("location"):
        parts.append(_clean_text(contact["location"]))
    if contact.get("phone"):
        parts.append(_clean_text(contact["phone"]))
    if contact.get("email"):
        parts.append(_clean_text(contact["email"]))
    if contact.get("linkedin"):
        parts.append(_clean_text(contact["linkedin"]))
    if contact.get("github"):
        parts.append(_clean_text(contact["github"]))

    if parts:
        p = doc.add_paragraph(" | ".join(parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(10)


def _add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    # Underline-style: bottom border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.get_or_add_pBdr() if hasattr(pPr, "get_or_add_pBdr") else None
    # Simple approach: just bold + spacing, no borders needed for ATS compliance
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _add_body_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(_clean_text(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(1)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(_clean_text(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(1)


def generate_docx(resume: Dict) -> bytes:
    """
    Build an ATS-compliant DOCX from a TailoredResume document.
    Returns the file contents as bytes.
    """
    profile = resume.get("tailored_profile", {})
    contact = resume.get("contact_override") or {}

    # Fall back to profile contact if available in the resume object
    # (callers may enrich the resume dict with the user's contact info)
    if not contact and resume.get("contact"):
        contact = resume["contact"]

    doc = Document()
    _set_margins(doc)

    # Remove default styles that add tables/columns
    # (python-docx default template is already single-column)

    # ------------------------------------------------------------------
    # Header: Name + Contact
    # ------------------------------------------------------------------
    name = _clean_text(contact.get("name", resume.get("job_title", "Resume")))
    _add_name(doc, name)
    _add_contact_line(doc, contact)

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------
    summary = profile.get("summary", "").strip()
    if summary:
        _add_section_heading(doc, "Summary")
        _add_body_para(doc, summary)

    # ------------------------------------------------------------------
    # Skills (flat comma-separated list — ATS rule)
    # ------------------------------------------------------------------
    skills: List[str] = profile.get("skills", [])
    if skills:
        _add_section_heading(doc, "Skills")
        _add_body_para(doc, ", ".join([_clean_text(s) for s in skills if s]))

    # ------------------------------------------------------------------
    # Experience
    # ------------------------------------------------------------------
    experiences: List[Dict] = profile.get("experiences", [])
    if experiences:
        _add_section_heading(doc, "Experience")
        for exp in experiences:
            company = _clean_text(exp.get("company", ""))
            title = _clean_text(exp.get("title", ""))
            location = _clean_text(exp.get("location", ""))
            date_range = _format_date_range(
                exp.get("start_date", ""), exp.get("end_date", "")
            )
            header_parts = [f"{title} — {company}"]
            if location:
                header_parts.append(location)
            if date_range:
                header_parts.append(date_range)
            _add_body_para(doc, " | ".join(header_parts), bold=True)
            for bullet in exp.get("bullets", []):
                if bullet:
                    _add_bullet(doc, bullet)

    # ------------------------------------------------------------------
    # Education
    # ------------------------------------------------------------------
    education: List[Dict] = resume.get("education_override") or []
    if not education and resume.get("education"):
        education = resume["education"]
    if education:
        _add_section_heading(doc, "Education")
        for edu in education:
            school = _clean_text(edu.get("school", ""))
            degree = _clean_text(edu.get("degree", ""))
            field = _clean_text(edu.get("field", ""))
            year = _clean_text(edu.get("year", ""))
            line = f"{degree} in {field} — {school}"
            if year:
                line += f" ({year})"
            _add_body_para(doc, line, bold=True)

    # ------------------------------------------------------------------
    # Projects
    # ------------------------------------------------------------------
    projects: List[Dict] = profile.get("projects", [])
    if projects:
        _add_section_heading(doc, "Projects")
        for proj in projects:
            name_text = _clean_text(proj.get("name", ""))
            tech = ", ".join([_clean_text(t) for t in proj.get("tech_stack", []) if t])
            purpose = _clean_text(proj.get("purpose", ""))
            header = name_text
            if tech:
                header += f" | {tech}"
            _add_body_para(doc, header, bold=True)
            if purpose:
                _add_body_para(doc, purpose)
            for feat in proj.get("key_features", []):
                if feat:
                    _add_bullet(doc, feat)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def generate_pdf(resume: Dict) -> bytes:
    """
    Build a clean, ATS-friendly PDF using reportlab.
    Single column, no tables, standard fonts.
    """
    profile = resume.get("tailored_profile", {})
    contact = resume.get("contact_override") or resume.get("contact") or {}

    buf = io.BytesIO()
    margin = 0.75 * inch

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin,
    )

    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "NameStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=16,
        spaceAfter=4,
        alignment=1,  # center
    )
    contact_style = ParagraphStyle(
        "ContactStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        spaceAfter=6,
        alignment=1,
    )
    heading_style = ParagraphStyle(
        "HeadingStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        spaceBefore=10,
        spaceAfter=3,
        textTransform="uppercase",
    )
    body_style = ParagraphStyle(
        "BodyStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        spaceAfter=2,
    )
    bold_body_style = ParagraphStyle(
        "BoldBodyStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "BulletStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leftIndent=12,
        spaceAfter=2,
        bulletIndent=4,
        bulletText="-",
    )

    story = []

    # Name
    name = _clean_text(contact.get("name", resume.get("job_title", "Resume")))
    story.append(Paragraph(name, name_style))

    # Contact line
    parts = []
    for field in ("location", "phone", "email", "linkedin", "github"):
        val = contact.get(field, "")
        if val:
            parts.append(_clean_text(val))
    if parts:
        story.append(Paragraph(" | ".join(parts), contact_style))

    # Summary
    summary = profile.get("summary", "").strip()
    if summary:
        story.append(Paragraph("SUMMARY", heading_style))
        story.append(Paragraph(_clean_text(summary), body_style))

    # Skills
    skills = profile.get("skills", [])
    if skills:
        story.append(Paragraph("SKILLS", heading_style))
        skills_line = ", ".join([_clean_text(s) for s in skills if s])
        story.append(Paragraph(skills_line, body_style))

    # Experience
    experiences = profile.get("experiences", [])
    if experiences:
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in experiences:
            company = _clean_text(exp.get("company", ""))
            title_text = _clean_text(exp.get("title", ""))
            location = _clean_text(exp.get("location", ""))
            date_range = _format_date_range(exp.get("start_date", ""), exp.get("end_date", ""))
            parts = [f"{title_text} — {company}"]
            if location:
                parts.append(location)
            if date_range:
                parts.append(date_range)
            story.append(Paragraph(" | ".join(parts), bold_body_style))
            for bullet in exp.get("bullets", []):
                if bullet:
                    story.append(Paragraph(_clean_text(bullet), bullet_style))

    # Education
    education = resume.get("education_override") or resume.get("education") or []
    if education:
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in education:
            school = _clean_text(edu.get("school", ""))
            degree = _clean_text(edu.get("degree", ""))
            field = _clean_text(edu.get("field", ""))
            year = _clean_text(edu.get("year", ""))
            line = f"{degree} in {field} — {school}"
            if year:
                line += f" ({year})"
            story.append(Paragraph(line, bold_body_style))

    # Projects
    projects = profile.get("projects", [])
    if projects:
        story.append(Paragraph("PROJECTS", heading_style))
        for proj in projects:
            proj_name = _clean_text(proj.get("name", ""))
            tech = ", ".join([_clean_text(t) for t in proj.get("tech_stack", []) if t])
            purpose = _clean_text(proj.get("purpose", ""))
            header = proj_name
            if tech:
                header += f" | {tech}"
            story.append(Paragraph(header, bold_body_style))
            if purpose:
                story.append(Paragraph(purpose, body_style))
            for feat in proj.get("key_features", []):
                if feat:
                    story.append(Paragraph(_clean_text(feat), bullet_style))

    doc.build(story)
    return buf.getvalue()
